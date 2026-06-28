"""
可复用文档组件 — 封面、目录、代码块、表格、Mermaid图插入等。

所有内容生成函数均通过 config 参数接收学校和课程信息，
绝不包含任何硬编码的个人/学校信息。
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dataclasses import dataclass, field
from typing import List, Optional, Dict
import os

from .styles import setup_styles, setup_page, _set_font, StyleConfig


# ═══════════════════════════════════════════════════════════════
# 项目元数据
# ═══════════════════════════════════════════════════════════════

@dataclass
class ProjectMeta:
    """项目/课程元数据 — 完全参数化，无默认个人信息。

    Examples:
        meta = ProjectMeta(
            course_name="综合设计项目",
            college="XX学院",
            semester="2025-2026 学年 2 学期",
            project_name="XXX项目",
            advisor="张教授",
            students=[{"name": "张三", "id": "20240001"}],
        )
    """
    course_name: str = ""
    college: str = ""
    semester: str = ""
    project_name: str = ""
    advisor: str = ""
    students: List[Dict[str, str]] = field(default_factory=list)
    abstract_text: str = ""
    keywords_text: str = ""


# ═══════════════════════════════════════════════════════════════
# 文档初始化
# ═══════════════════════════════════════════════════════════════

def new_document(style_config: StyleConfig = None):
    """创建并初始化新 Document，含样式。

    Args:
        style_config: 样式配置，None 则使用默认

    Returns:
        Document 对象
    """
    doc = Document()
    if style_config is None:
        style_config = StyleConfig()
    setup_styles(doc, style_config)
    return doc


# ═══════════════════════════════════════════════════════════════
# 封面 + 摘要 + 目录
# ═══════════════════════════════════════════════════════════════

def create_cover(doc: Document, meta: ProjectMeta,
                 style_config: StyleConfig = None):
    """生成封面页 + 摘要页。

    Args:
        doc: Document 对象
        meta: 项目元数据
        style_config: 样式配置
    """
    if style_config is None:
        style_config = StyleConfig()
    setup_page(doc, "cover", style_config)

    for _ in range(5):
        doc.add_paragraph("", 'Normal')

    # 课程名称
    p = doc.add_paragraph(f"课程名称：{meta.course_name}", 'CoverInfo')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"所在系别：{meta.college}", 'CoverInfo')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"学    期：{meta.semester}", 'CoverInfo')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"课题名称：{meta.project_name}", 'CoverInfo')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"指导教师：{meta.advisor}", 'CoverInfo')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("", 'Normal')

    p = doc.add_paragraph("学生信息：", 'CoverInfo')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    students = meta.students if meta.students else [
        {"name": "", "id": ""} for _ in range(5)
    ]

    table = doc.add_table(rows=len(students) + 1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, text in enumerate(["序号", "学号", "姓名"]):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_font(run, style_config.en_font, style_config.cn_font_body,
                  style_config.size_body, bold=True)

    for row_idx, stu in enumerate(students):
        for col_idx, key in enumerate(["index", "id", "name"]):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if key == "index":
                text = str(row_idx + 1)
            else:
                text = stu.get(key, "")
            run = p.add_run(text)
            _set_font(run, style_config.en_font, style_config.cn_font_body,
                      style_config.size_body)

    doc.add_page_break()

    # 摘要
    p = doc.add_paragraph("摘  要", 'AbstractTitle')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if meta.abstract_text:
        p = doc.add_paragraph(meta.abstract_text, 'AbstractBody')
    else:
        p = doc.add_paragraph("（请填写摘要内容）", 'AbstractBody')

    if meta.keywords_text:
        p = doc.add_paragraph(f"关键词：{meta.keywords_text}", 'Keywords')
    else:
        p = doc.add_paragraph("关键词：", 'Keywords')

    doc.add_page_break()
    return doc


def insert_toc_field(doc: Document):
    """插入 Word 目录域代码。在 Word 中右键 → 更新域即可生成目录。"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = paragraph.add_run("目  录")
    _set_font(run_title, 'Times New Roman', '黑体', 15, bold=True)

    p = doc.add_paragraph()

    run = p.add_run()
    fldChar = run._r.makeelement(qn('w:fldChar'),
                                  {qn('w:fldCharType'): 'begin'})
    run._r.append(fldChar)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = run3._r.makeelement(qn('w:fldChar'),
                                    {qn('w:fldCharType'): 'separate'})
    run3._r.append(fldChar2)

    run4 = p.add_run('【请在 Microsoft Word 中右键此处 → 更新域 → 更新整个目录】')
    _set_font(run4, 'Times New Roman', '宋体', 12, bold=False)
    run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    run5 = p.add_run()
    fldChar3 = run5._r.makeelement(qn('w:fldChar'),
                                    {qn('w:fldCharType'): 'end'})
    run5._r.append(fldChar3)

    doc.add_page_break()
    return doc


# ═══════════════════════════════════════════════════════════════
# 段落 / 文本
# ═══════════════════════════════════════════════════════════════

def add_heading(doc: Document, text: str, level: int = 1):
    """添加标题 (Heading 1/2/3)。"""
    return doc.add_heading(text, level=level)


def add_body_para(doc: Document, text: str):
    """添加正文段落 (首行缩进, 1.5倍行距)。"""
    return doc.add_paragraph(text, 'BodyText12')


def add_bold_para(doc: Document, text: str):
    """添加粗体强调段落。"""
    p = doc.add_paragraph("", 'BodyText12')
    run = p.add_run(text)
    run.bold = True
    return p


def add_inline_emphasized_para(doc: Document, prefix: str = "",
                               bold_part: str = "", suffix: str = ""):
    """添加 前缀 + 粗体关键字 + 后缀 的混合段落。"""
    p = doc.add_paragraph("", 'BodyText12')
    if prefix:
        p.add_run(prefix)
    run_bold = p.add_run(bold_part)
    run_bold.bold = True
    if suffix:
        p.add_run(suffix)
    return p


def add_list_item(doc: Document, text: str, level: int = 0):
    """添加列表项 (• 开头)。"""
    p = doc.add_paragraph("", 'BodyText12')
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.clear()
    p.add_run(f"• {text}")
    return p


def add_numbered_item(doc: Document, number: int, text: str,
                      level: int = 0):
    """添加编号列表项。"""
    p = doc.add_paragraph("", 'BodyText12')
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.add_run(f"{number}. {text}")
    return p


# ═══════════════════════════════════════════════════════════════
# 代码块 (语法高亮)
# ═══════════════════════════════════════════════════════════════

def add_code_block(doc: Document, code_text: str = "",
                   language: str = "cpp",
                   style_config: StyleConfig = None,
                   caption: str = "",
                   file: str = "",
                   start_line: int = 0,
                   end_line: int = 0,
                   highlight_lines: list = None):
    """带背景色+行号+语法着色的代码块。

    Args:
        doc: Document
        code_text: 代码字符串 (与 file 二选一)
        language: "cpp"|"c"|"python"|"py"|"llvm"|"ir"|
                  "riscv"|"rv"|"arm"|"aarch64"|"x86"|"asm"|"bash"|"sh"|""
        style_config: StyleConfig
        caption: 代码块标题 (如 "代码4-1 核心算法核心算法")
        file: 从文件读取代码 (与 code_text 二选一)
        start_line: 起始行 (1-based, 配合 file 使用)
        end_line: 结束行 (1-based, 配合 file 使用; 0=到末尾)
        highlight_lines: 高亮行号列表, 如 [12, 13, 14]
    """
    import re
    if style_config is None:
        style_config = StyleConfig()

    # 从文件读取
    if file and os.path.exists(file):
        with open(file, 'r', encoding='utf-8', errors='replace') as f:
            all_lines = f.readlines()
        if end_line == 0:
            end_line = len(all_lines)
        if start_line > 0:
            all_lines = all_lines[start_line - 1:end_line]
        code_text = ''.join(all_lines)
        # 从扩展名推断语言
        if language == "cpp":
            ext = os.path.splitext(file)[1].lower()
            ext_map = {'.py': 'python', '.c': 'c', '.h': 'c', '.hpp': 'cpp',
                       '.cpp': 'cpp', '.cc': 'cpp', '.ll': 'llvm',
                       '.s': 'riscv', '.S': 'riscv', '.asm': 'x86',
                       '.sh': 'bash'}
            language = ext_map.get(ext, language)

    if not code_text:
        code_text = ""

    lines = code_text.split('\n')
    nw = len(str(len(lines)))

    # 选择关键词集
    lang_lower = language.lower()
    if lang_lower in ("cpp", "c"):
        KWDS = _CPP_KEYWORDS
    elif lang_lower in ("py", "python"):
        KWDS = _PY_KEYWORDS
    elif lang_lower in ("llvm", "ir"):
        KWDS = _CPP_KEYWORDS | _中间表示_KEYWORDS
    elif lang_lower in ("riscv", "rv"):
        KWDS = _RISCV_KEYWORDS
    elif lang_lower in ("arm", "aarch64"):
        KWDS = _目标平台_KEYWORDS
    elif lang_lower in ("x86", "asm"):
        KWDS = _X86_KEYWORDS
    elif lang_lower in ("bash", "sh"):
        KWDS = _BASH_KEYWORDS
    else:
        KWDS = _CPP_KEYWORDS | _PY_KEYWORDS | _中间表示_KEYWORDS

    TYPES = _TYPE_KEYWORDS
    if highlight_lines is None:
        highlight_lines = []

    # 标题
    if caption:
        cap = doc.add_paragraph(caption, 'FigCaption')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    spacer = doc.add_paragraph("", 'Normal')
    spacer.paragraph_format.space_before = Pt(6)
    spacer.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].clear()

    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), style_config.code_bg)
    shading.set(qn('w:color'), 'auto')
    tcPr.append(shading)

    TOKEN_RE = re.compile(
        r'(//[^\n]*|/\*.*?\*/|#.*|;.*)'      # comments
        r'|("(?:[^"\\]|\\.)*")'              # double-quoted
        r"|('(?:[^'\\]|\\.)*')"              # single-quoted
        r'|(\b0[xX][0-9a-fA-F]+\b|\b\d+\.?\d*(?:[eE][+-]?\d+)?\b)'
        r'|([a-zA-Z_]\w*)'
        r'|([{}()\[\];,:<>=+\-*/%&|^!~.@#\\]+)'
        r'|(\s+)', re.DOTALL
    )

    def colorize_line(line):
        tokens = []
        for m in TOKEN_RE.finditer(line):
            txt = m.group(0)
            if m.group(1):
                tokens.append((txt, style_config.code_comment))
            elif m.group(2) or m.group(3):
                tokens.append((txt, style_config.code_string))
            elif m.group(4):
                tokens.append((txt, style_config.code_number))
            elif m.group(5):
                if txt in KWDS:
                    tokens.append((txt, style_config.code_keyword))
                elif txt in TYPES:
                    tokens.append((txt, style_config.code_type))
                elif txt.startswith('x') and txt[1:].isdigit():
                    # 目标平台 寄存器 x0-x30
                    tokens.append((txt, RGBColor(0xC0, 0x00, 0x00)))
                elif txt.startswith('v') and txt[1:].isdigit():
                    # 目标平台 SIMD 寄存器 v0-v31
                    tokens.append((txt, RGBColor(0xC0, 0x00, 0x00)))
                else:
                    tokens.append((txt, style_config.color_black))
            elif m.group(6):
                if txt.startswith('@') or txt.startswith('%'):
                    tokens.append((txt, RGBColor(0xC0, 0x00, 0x00)))
                else:
                    tokens.append((txt, RGBColor(0x80, 0x80, 0x80)))
            else:
                tokens.append((txt, style_config.color_black))
        return tokens

    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05

        # 行高亮背景
        line_num = i + 1
        if line_num in highlight_lines:
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), 'FFF3CD')  # 淡黄
            pPr.append(shd)

        rn = p.add_run(f" {(line_num):>{nw}} │ ")
        rn.font.name = style_config.code_font
        rn.font.size = Pt(8)
        rn.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        colored = (colorize_line(line) if line.strip()
                   else [(line if line else ' ', style_config.color_black)])
        for txt, color in colored:
            rc = p.add_run(txt)
            rc.font.name = style_config.code_font
            rc.font.size = Pt(style_config.size_code)
            rc.font.color.rgb = color

    spacer2 = doc.add_paragraph("", 'Normal')
    spacer2.paragraph_format.space_before = Pt(0)
    spacer2.paragraph_format.space_after = Pt(6)
    return doc


_CPP_KEYWORDS = {
    'int', 'float', 'double', 'void', 'const', 'if', 'else', 'while',
    'break', 'continue', 'return', 'for', 'do', 'switch', 'case',
    'default', 'struct', 'class', 'enum', 'typedef', 'sizeof',
    'static', 'extern', 'inline', 'volatile', 'unsigned', 'signed',
    'char', 'short', 'long', 'bool', 'true', 'false', 'nullptr',
    'NULL', 'public', 'protected', 'private', 'virtual', 'override',
    'final', 'template', 'typename', 'namespace', 'using', 'new',
    'delete', 'auto', 'register', 'include', 'pragma', 'define',
    'ifdef', 'ifndef', 'endif', 'std', 'unique_ptr', 'shared_ptr',
    'weak_ptr', 'make_shared', 'make_unique', 'vector', 'map',
    'unordered_map', 'string', 'constexpr', 'noexcept', 'explicit',
}

_PY_KEYWORDS = {
    'def', 'class', 'import', 'from', 'as', 'if', 'elif', 'else',
    'while', 'for', 'in', 'break', 'continue', 'return', 'yield',
    'try', 'except', 'finally', 'raise', 'with', 'pass', 'lambda',
    'True', 'False', 'None', 'and', 'or', 'not', 'is', 'self',
    'print', 'range', 'len', 'str', 'int', 'float', 'list', 'dict',
    'set', 'tuple', 'type', 'isinstance', 'super', '__init__',
}

_中间表示_KEYWORDS = {
    'define', 'declare', 'target', 'datalayout', 'triple', 'module',
    'global', 'constant', 'private', 'internal', 'available_externally',
    'linkonce', 'noreturn', 'nounwind', 'noinline', 'call', 'ret',
    'br', 'alloca', 'store', 'load', 'getelementptr', 'gep',
    'add', 'sub', 'mul', 'sdiv', 'udiv', 'srem', 'urem',
    'fadd', 'fsub', 'fmul', 'fdiv', 'frem',
    'shl', 'lshr', 'ashr', 'and', 'or', 'xor',
    'icmp', 'fcmp', 'phi', 'select',
    'trunc', 'zext', 'sext', 'fptosi', 'sitofp', 'bitcast',
    'uitofp', 'fptoui',
}

_TYPE_KEYWORDS = {
    'i1', 'i8', 'i16', 'i32', 'i64', 'i128',
    'half', 'bfloat', 'float', 'double', 'x86_fp80',
    'fp128', 'void', 'label', 'metadata', 'token', 'ptr', 'opaque',
    'Value', 'User', 'Use', 'Instruction', 'BasicBlock',
    'Function', 'Module', 'IRType', 'TypeSystem',
    'PreservationStatus', 'PassID', 'PipelineManager',
}

_RISCV_KEYWORDS = {
    # 目标平台 instructions (RV64G)
    'add', 'addi', 'addiw', 'addw', 'sub', 'subw',
    'mul', 'mulw', 'div', 'divw', 'divu', 'divuw',
    'rem', 'remw', 'remu', 'remuw',
    'and', 'andi', 'or', 'ori', 'xor', 'xori',
    'sll', 'slli', 'sllw', 'slliw', 'srl', 'srli', 'srlw', 'srliw',
    'sra', 'srai', 'sraw', 'sraiw',
    'slt', 'slti', 'sltu', 'sltiu',
    'lb', 'lh', 'lw', 'ld', 'lbu', 'lhu', 'lwu',
    'sb', 'sh', 'sw', 'sd',
    'beq', 'bne', 'blt', 'bge', 'bltu', 'bgeu',
    'jal', 'jalr', 'j', 'jr', 'ret',
    'lui', 'auipc',
    'ecall', 'ebreak', 'fence', 'fence.i',
    # Float
    'fadd', 'fsub', 'fmul', 'fdiv', 'fsqrt',
    'fadd.s', 'fsub.s', 'fmul.s', 'fdiv.s',
    'fcvt.s.d', 'fcvt.d.s', 'fcvt.w.s', 'fcvt.s.w',
    'feq', 'flt', 'fle',
    'fld', 'flw', 'fsd', 'fsw',
    'fmv', 'fmv.s', 'fmv.d', 'fmv.w.x', 'fmv.x.w',
    # Pseudo
    'mv', 'li', 'la', 'nop', 'call', 'tail',
    # CSRs
    'csrr', 'csrw', 'csrs', 'csrc',
    # Registers
    'zero', 'ra', 'sp', 'gp', 'tp',
    't0', 't1', 't2', 't3', 't4', 't5', 't6',
    's0', 's1', 's2', 's3', 's4', 's5', 's6', 's7', 's8', 's9', 's10', 's11',
    'a0', 'a1', 'a2', 'a3', 'a4', 'a5', 'a6', 'a7',
    'ft0', 'ft1', 'ft2', 'ft3', 'ft4', 'ft5', 'ft6', 'ft7', 'ft8', 'ft9', 'ft10', 'ft11',
    'fs0', 'fs1', 'fs2', 'fs3', 'fs4', 'fs5', 'fs6', 'fs7', 'fs8', 'fs9', 'fs10', 'fs11',
    'fa0', 'fa1', 'fa2', 'fa3', 'fa4', 'fa5', 'fa6', 'fa7',
    # Directives
    '.text', '.data', '.bss', '.section', '.global', '.globl',
    '.align', '.byte', '.half', '.word', '.dword', '.asciz', '.string',
    '.equ', '.macro', '.endm', '.type', '.size', '.file', '.loc',
}

_目标平台_KEYWORDS = {
    # AArch64 instructions
    'add', 'adds', 'sub', 'subs', 'mul', 'madd', 'msub',
    'sdiv', 'udiv', 'and', 'ands', 'orr', 'eor', 'eon',
    'lsl', 'lsr', 'asr', 'ror',
    'mov', 'movk', 'movn', 'movz',
    'ldr', 'ldrb', 'ldrh', 'ldrsw', 'ldp', 'ldnp',
    'str', 'strb', 'strh', 'stp', 'stnp',
    'cbz', 'cbnz', 'tbz', 'tbnz',
    'b', 'bl', 'blr', 'br', 'ret',
    'cmp', 'cmn', 'tst',
    'csel', 'csinc', 'csinv', 'csneg',
    'adr', 'adrp',
    'fadd', 'fsub', 'fmul', 'fdiv', 'fmadd', 'fmsub',
    'fcmp', 'fccmp', 'fcsel',
    'fcvt', 'scvtf', 'ucvtf', 'fcvtzs', 'fcvtzu',
    'ld1', 'ld2', 'ld3', 'ld4', 'st1', 'st2', 'st3', 'st4',
    'fmov', 'fabs', 'fneg', 'fsqrt',
    'svc', 'hlt', 'nop', 'msr', 'mrs',
    # Conditions
    'eq', 'ne', 'cs', 'hs', 'cc', 'lo', 'mi', 'pl',
    'vs', 'vc', 'hi', 'ls', 'ge', 'lt', 'gt', 'le', 'al',
    # Directives
    '.text', '.data', '.bss', '.section', '.global', '.globl',
    '.align', '.byte', '.hword', '.word', '.quad', '.asciz', '.string',
    '.equ', '.macro', '.endm', '.type', '.size', '.cfi_startproc', '.cfi_endproc',
}

_X86_KEYWORDS = {
    'mov', 'movsx', 'movzx', 'push', 'pop',
    'add', 'sub', 'mul', 'imul', 'div', 'idiv',
    'and', 'or', 'xor', 'not', 'neg',
    'shl', 'shr', 'sar', 'rol', 'ror',
    'lea', 'call', 'ret', 'jmp', 'je', 'jne', 'jz', 'jnz',
    'jl', 'jle', 'jg', 'jge', 'jb', 'jbe', 'ja', 'jae',
    'cmp', 'test',
    'inc', 'dec', 'nop', 'int',
    'rax', 'rbx', 'rcx', 'rdx', 'rsi', 'rdi', 'rbp', 'rsp',
    'r8', 'r9', 'r10', 'r11', 'r12', 'r13', 'r14', 'r15',
    'eax', 'ebx', 'ecx', 'edx', 'esi', 'edi', 'ebp', 'esp',
    'xmm0', 'xmm1', 'xmm2', 'xmm3', 'xmm4', 'xmm5',
    'xmm6', 'xmm7', 'xmm8', 'xmm9', 'xmm10', 'xmm11',
    'xmm12', 'xmm13', 'xmm14', 'xmm15',
    '.text', '.data', '.bss', '.section', '.global', '.globl',
    '.align', '.byte', '.word', '.long', '.quad', '.asciz', '.string',
}

_BASH_KEYWORDS = {
    'if', 'then', 'else', 'elif', 'fi', 'for', 'while', 'do', 'done',
    'case', 'esac', 'in', 'function', 'return', 'exit',
    'export', 'local', 'readonly', 'unset',
    'echo', 'printf', 'cd', 'ls', 'mkdir', 'rm', 'cp', 'mv',
    'grep', 'sed', 'awk', 'cat', 'head', 'tail', 'find', 'xargs',
    'source', '.', '&&', '||', '|',
}


# ═══════════════════════════════════════════════════════════════
# 图表
# ═══════════════════════════════════════════════════════════════

def add_mermaid_image(doc: Document, image_path: str,
                      caption: str = "", width_cm: float = None):
    """插入 Mermaid 渲染后的图片，自动限高≤18cm。

    Args:
        doc: Document
        image_path: PNG 文件路径 (绝对或相对)
        caption: 图题，如 "图 3-1 系统总体架构"
        width_cm: 指定宽度(cm)，默认14cm
    """
    from PIL import Image as PILImage

    MAX_HEIGHT_CM = 18.0

    if not os.path.exists(image_path):
        p = doc.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'[图片缺失: {os.path.basename(image_path)}]')
        _set_font(run, 'Courier New', '宋体', 10)
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        if caption:
            doc.add_paragraph(caption, 'FigCaption')
        return doc

    pil_img = PILImage.open(image_path)
    img_w_px, img_h_px = pil_img.size
    aspect = img_w_px / img_h_px

    if width_cm is None:
        target_w = 14.0
    else:
        target_w = width_cm

    implied_h = target_w / aspect
    if implied_h > MAX_HEIGHT_CM:
        target_w = MAX_HEIGHT_CM * aspect

    p = doc.add_paragraph("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run()
    run.add_picture(image_path, width=Cm(target_w))

    if caption:
        cap = doc.add_paragraph(caption, 'FigCaption')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def add_figure(doc: Document, img_path: str, caption: str = "",
               width_cm: float = 14.0):
    """插入任意 PNG/JPG 图片。"""
    p = doc.add_paragraph("")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

    if os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    else:
        run = p.add_run(f'[图片缺失: {os.path.basename(img_path)}]')
        _set_font(run, 'Courier New', '宋体', 10)
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    if caption:
        cap = doc.add_paragraph(caption, 'FigCaption')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


# ═══════════════════════════════════════════════════════════════
# 表格
# ═══════════════════════════════════════════════════════════════

def add_table(doc: Document, headers: List[str], rows: List[List[str]],
              caption: str = "", col_widths: List = None,
              style_config: StyleConfig = None):
    """插入三线表风格的表格。

    Args:
        doc: Document
        headers: ["列1", "列2", ...]
        rows: [["a","b"], ["c","d"], ...]
        caption: 表题
        col_widths: [Cm(3), Cm(10), ...] (可选)
        style_config: StyleConfig
    """
    if style_config is None:
        style_config = StyleConfig()

    if caption:
        cap = doc.add_paragraph(caption, 'FigCaption')
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = width

    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        _set_font(run, style_config.en_font, style_config.cn_font_body,
                  style_config.size_body, bold=True)

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_text))
            _set_font(run, style_config.en_font, style_config.cn_font_body,
                      style_config.size_body, bold=False)

    doc.add_paragraph("", 'Normal')
    return doc


# ═══════════════════════════════════════════════════════════════
# 章节管理
# ═══════════════════════════════════════════════════════════════

def add_chapter_start(doc: Document, chapter_title: str,
                      chapter_number: int):
    """开始新章节 (换页 + Heading 1)。"""
    doc.add_page_break()
    add_heading(doc, f"第{chapter_number}章 {chapter_title}", level=1)
    return doc


# ═══════════════════════════════════════════════════════════════
# 参考文献
# ═══════════════════════════════════════════════════════════════

def add_references(doc: Document, refs: List[str]):
    """添加参考文献列表。

    Args:
        doc: Document
        refs: 参考文献字符串列表，如 ["[1] 作者. 标题[M]. 出版社, 2020."]
    """
    add_heading(doc, "参考文献", level=1)
    for ref in refs:
        p = doc.add_paragraph(ref, 'BodyText12')
        p.paragraph_format.first_line_indent = Cm(0)
    return doc


# ═══════════════════════════════════════════════════════════════
# 页眉/页脚/页码
# ═══════════════════════════════════════════════════════════════

def add_page_numbers(doc: Document, position: str = "bottom_center",
                     start_at: int = 1, style_config: StyleConfig = None):
    """添加页码。

    Args:
        doc: Document
        position: "bottom_center" | "bottom_right" | "top_center" | "top_right"
        start_at: 起始页码
        style_config: StyleConfig
    """
    if style_config is None:
        style_config = StyleConfig()

    section = doc.sections[-1]
    footer = section.footer
    footer.is_linked_to_previous = False

    fp = footer.paragraphs[0]
    if 'center' in position:
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # PageNumber field
    run = fp.add_run()
    fldChar1 = run._r.makeelement(qn('w:fldChar'),
                                   {qn('w:fldCharType'): 'begin'})
    run._r.append(fldChar1)

    run2 = fp.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._r.append(instrText)

    run3 = fp.add_run()
    fldChar2 = run3._r.makeelement(qn('w:fldChar'),
                                    {qn('w:fldCharType'): 'separate'})
    run3._r.append(fldChar2)

    run4 = fp.add_run(str(start_at))
    _set_font(run4, style_config.en_font, style_config.cn_font_body,
              style_config.size_body)

    run5 = fp.add_run()
    fldChar3 = run5._r.makeelement(qn('w:fldChar'),
                                    {qn('w:fldCharType'): 'end'})
    run5._r.append(fldChar3)

    # 设置起始页码
    if start_at != 1:
        sectPr = section._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), str(start_at))
        sectPr.append(pgNumType)

    return doc


def add_header_footer(doc: Document,
                      header_text: str = "",
                      show_page_number: bool = True,
                      style_config: StyleConfig = None):
    """统一设置页眉和页脚。

    Args:
        doc: Document
        header_text: 页眉文字 (为空则不设页眉)
        show_page_number: 是否显示页码
        style_config: StyleConfig
    """
    if style_config is None:
        style_config = StyleConfig()

    section = doc.sections[-1]

    if header_text:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(header_text)
        _set_font(run, style_config.en_font, style_config.cn_font_body,
                  style_config.size_header)

    if show_page_number:
        add_page_numbers(doc, style_config=style_config)

    return doc
