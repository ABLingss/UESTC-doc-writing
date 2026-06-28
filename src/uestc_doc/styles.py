"""
样式定义模块 — 针对学术报告/论文模板优化。
所有参数可覆盖，无个人信息硬编码。

Usage:
    from uestc_doc.styles import setup_styles, setup_page, StyleConfig
    doc = Document()
    config = StyleConfig(header_text="学术报告")
    setup_styles(doc, config)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from dataclasses import dataclass, field


# ── 格式常量 ────────────────────────────────────────
A4_WIDTH = Cm(21.0)
A4_HEIGHT = Cm(29.7)


@dataclass
class StyleConfig:
    """可覆盖的样式配置 — 适配不同院系/课程模板。

    所有字段均有默认值（电子科大学术规范），可通过构造参数覆盖。
    """
    # ── 页边距 ──
    margin_top_body: object = Cm(2.88)
    margin_bottom_body: object = Cm(1.97)
    margin_left_body: object = Cm(2.36)
    margin_right_body: object = Cm(2.36)
    margin_top_cover: object = Cm(0.56)
    margin_bottom_cover: object = Cm(0.56)

    # ── 页眉 ──
    header_text: str = "学术报告"

    # ── 字体族 ──
    en_font: str = "Times New Roman"
    cn_font_body: str = "宋体"
    cn_font_heading: str = "黑体"
    code_font: str = "Courier New"

    # ── 字号 (pt) ──
    size_body: float = 12.0
    size_h1: float = 15.0
    size_h2: float = 12.0
    size_h3: float = 10.5
    size_cover_title: float = 28.0
    size_cover_info: float = 15.0
    size_code: float = 9.0
    size_toc: float = 11.0
    size_header: float = 9.0

    # ── 行距 ──
    line_spacing_body: float = 1.5
    line_spacing_toc: float = 1.08

    # ── 缩进 ──
    first_line_indent: object = Cm(0.74)
    toc_level2_indent: object = Cm(0.78)
    toc_level3_indent: object = Cm(1.56)

    # ── 颜色 ──
    color_black: object = RGBColor(0x00, 0x00, 0x00)
    color_toc: object = RGBColor(0x37, 0x60, 0x92)

    # ── 代码块颜色 ──
    code_bg: str = "F8F8F8"
    code_keyword: object = field(default_factory=lambda: RGBColor(0x00, 0x00, 0xFF))
    code_comment: object = field(default_factory=lambda: RGBColor(0x00, 0x80, 0x00))
    code_string: object = field(default_factory=lambda: RGBColor(0xA0, 0x50, 0x00))
    code_number: object = field(default_factory=lambda: RGBColor(0x80, 0x00, 0x80))
    code_type: object = field(default_factory=lambda: RGBColor(0x00, 0x80, 0x80))


# ── 内部工具函数 ────────────────────────────────────

def _set_east_asian_font(run, font_name: str):
    """设置东亚字体 (中文回退字体)。"""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _set_font(run, en_font: str, cn_font: str, size_pt: float,
              bold: bool = False, color=None):
    """设置 run 的英/中字体、字号、粗细、颜色。"""
    if color is None:
        color = RGBColor(0, 0, 0)
    run.font.name = en_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color
    _set_east_asian_font(run, cn_font)


# ── 样式注册 ────────────────────────────────────────

def setup_styles(doc: Document, config: StyleConfig = None):
    """在 Document 上注册所有自定义样式。

    Args:
        doc: python-docx Document 对象
        config: StyleConfig 实例，为 None 则使用默认配置

    Returns:
        Document (链式调用)
    """
    if config is None:
        config = StyleConfig()

    # ── Normal / 正文 ──────────────────────────────
    style = doc.styles['Normal']
    style.font.name = config.en_font
    style.font.size = Pt(config.size_body)
    style.font.color.rgb = config.color_black
    style.paragraph_format.line_spacing = config.line_spacing_body
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), config.cn_font_body)

    # ── Heading 1 ──────────────────────────────────
    h1 = doc.styles['Heading 1']
    h1.font.name = config.en_font
    h1.font.size = Pt(config.size_h1)
    h1.font.bold = True
    h1.font.color.rgb = config.color_black
    h1.paragraph_format.line_spacing = config.line_spacing_body
    h1.paragraph_format.space_before = Pt(17)
    h1.paragraph_format.space_after = Pt(16)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.element.rPr.rFonts.set(qn('w:eastAsia'), config.cn_font_heading)

    # ── Heading 2 ──────────────────────────────────
    h2 = doc.styles['Heading 2']
    h2.font.name = config.en_font
    h2.font.size = Pt(config.size_h2)
    h2.font.bold = True
    h2.font.color.rgb = config.color_black
    h2.paragraph_format.line_spacing = config.line_spacing_body
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.element.rPr.rFonts.set(qn('w:eastAsia'), config.cn_font_heading)

    # ── Heading 3 ──────────────────────────────────
    h3 = doc.styles['Heading 3']
    h3.font.name = config.en_font
    h3.font.size = Pt(config.size_h3)
    h3.font.bold = True
    h3.font.color.rgb = config.color_black
    h3.paragraph_format.line_spacing = config.line_spacing_body
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.element.rPr.rFonts.set(qn('w:eastAsia'), config.cn_font_heading)

    # ── 封面标题 ──────────────────────────────────
    _reg_para_style(doc, 'CoverTitle', config.en_font, config.cn_font_heading,
                    config.size_cover_title, False, WD_ALIGN_PARAGRAPH.CENTER,
                    Pt(20), Pt(0), 1.12)

    # ── 封面信息行 ─────────────────────────────────
    _reg_para_style(doc, 'CoverInfo', config.en_font, config.cn_font_body,
                    config.size_cover_info, False, WD_ALIGN_PARAGRAPH.CENTER,
                    Pt(16), Pt(0), 1.0)

    # ── 摘要标题 ──────────────────────────────────
    _reg_para_style(doc, 'AbstractTitle', config.en_font, config.cn_font_heading,
                    config.size_h1, True, WD_ALIGN_PARAGRAPH.CENTER,
                    Pt(20), Pt(10), config.line_spacing_body)

    # ── 摘要正文 ──────────────────────────────────
    _reg_para_style(doc, 'AbstractBody', config.en_font, config.cn_font_body,
                    config.size_body, False, WD_ALIGN_PARAGRAPH.JUSTIFY,
                    Pt(0), Pt(0), config.line_spacing_body,
                    first_line_indent=config.first_line_indent)

    # ── 关键词 ────────────────────────────────────
    _reg_para_style(doc, 'Keywords', config.en_font, config.cn_font_body,
                    config.size_body, False, WD_ALIGN_PARAGRAPH.JUSTIFY,
                    Pt(0), Pt(0), config.line_spacing_body,
                    first_line_indent=config.first_line_indent)

    # ── 正文段落 ──────────────────────────────────
    _reg_para_style(doc, 'BodyText12', config.en_font, config.cn_font_body,
                    config.size_body, False, WD_ALIGN_PARAGRAPH.JUSTIFY,
                    Pt(0), Pt(0), config.line_spacing_body,
                    first_line_indent=config.first_line_indent)

    # ── 图题 / 表题 ────────────────────────────────
    _reg_para_style(doc, 'FigCaption', config.en_font, config.cn_font_heading,
                    config.size_body, False, WD_ALIGN_PARAGRAPH.CENTER,
                    Pt(6), Pt(12), 1.0)

    # ── Code Block ────────────────────────────────
    _reg_para_style(doc, 'CodeBlock', config.code_font, config.cn_font_body,
                    config.size_code, False, WD_ALIGN_PARAGRAPH.LEFT,
                    Pt(0), Pt(0), 1.0, first_line_indent=Cm(0))

    # ── TOC 样式 ──────────────────────────────────
    _reg_para_style(doc, 'toc1', config.en_font, config.cn_font_body,
                    config.size_toc, True, WD_ALIGN_PARAGRAPH.LEFT,
                    Pt(0), Pt(5), config.line_spacing_toc)

    _reg_para_style(doc, 'toc2', config.en_font, config.cn_font_body,
                    config.size_toc, False, WD_ALIGN_PARAGRAPH.LEFT,
                    Pt(0), Pt(3), config.line_spacing_toc,
                    left_indent=config.toc_level2_indent)

    _reg_para_style(doc, 'toc3', config.en_font, config.cn_font_body,
                    config.size_toc, False, WD_ALIGN_PARAGRAPH.LEFT,
                    Pt(0), Pt(2), config.line_spacing_toc,
                    left_indent=config.toc_level3_indent)

    # ── 表格 ──────────────────────────────────────
    _reg_para_style(doc, 'TableText', config.en_font, config.cn_font_body,
                    config.size_body, False, WD_ALIGN_PARAGRAPH.CENTER,
                    Pt(2), Pt(2), 1.0)
    _reg_para_style(doc, 'TableHeader', config.en_font, config.cn_font_body,
                    config.size_body, True, WD_ALIGN_PARAGRAPH.CENTER,
                    Pt(2), Pt(2), 1.0)

    return doc


def _reg_para_style(doc, name, en_font, cn_font, size, bold, alignment,
                    space_before, space_after, line_spacing,
                    first_line_indent=None, left_indent=None):
    """注册一个段落样式。"""
    try:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        style = doc.styles[name]
    style.font.name = en_font
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.alignment = alignment
    style.paragraph_format.space_before = space_before
    style.paragraph_format.space_after = space_after
    style.paragraph_format.line_spacing = line_spacing
    if first_line_indent is not None:
        style.paragraph_format.first_line_indent = first_line_indent
    if left_indent is not None:
        style.paragraph_format.left_indent = left_indent
    style.element.rPr.rFonts.set(qn('w:eastAsia'), cn_font)
    return style


def setup_page(doc: Document, section_type: str = "body",
               config: StyleConfig = None):
    """配置当前 section 的页面布局。

    Args:
        doc: Document 对象
        section_type: "cover" | "body" | "first_body"
        config: StyleConfig 实例
    """
    if config is None:
        config = StyleConfig()

    section = doc.sections[-1]
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT

    if section_type == "cover":
        section.top_margin = config.margin_top_cover
        section.bottom_margin = config.margin_bottom_cover
        section.left_margin = Cm(2.51)
        section.right_margin = Cm(2.51)
        section.different_first_page_header_footer = True
        section.header.distance = Cm(0)
        section.footer.distance = Cm(0)
    else:
        section.top_margin = config.margin_top_body
        section.bottom_margin = config.margin_bottom_body
        section.left_margin = config.margin_left_body
        section.right_margin = config.margin_right_body
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = hp.add_run(config.header_text)
        _set_font(run, config.en_font, config.cn_font_body,
                  config.size_header)

    return doc
