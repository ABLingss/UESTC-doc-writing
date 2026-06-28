"""
UESTC 综合设计报告 Builder。

组装封面+目录+各章节, 生成完整 .docx 文档。

报告结构:
- 中期报告 (3章 + 参考文献)
- 后期报告 (4章 + 参考文献 + 致谢)
"""

from enum import Enum
from docx import Document

from uestc_doc import (
    StyleConfig, ProjectMeta,
    new_document, create_cover, insert_toc_field,
    add_chapter_start, add_heading, add_body_para, add_bold_para,
    add_list_item, add_code_block, add_table, add_references,
)

from .sections import (
    # 中期专用
    write_design_solution,
    write_reasoning_analysis,
    write_implementation_progress,
    write_knowledge_skills,
    write_completion_and_plan,
    # 后期专用
    write_requirements_modeling,
    write_complex_problems,
    write_feasibility_study,
    write_execution_overview,
    write_core_completion,
    write_simulation_results,
    write_completion_evaluation,
    write_team_collaboration,
    # 共用
    write_problems_and_solutions,
)


class ReportType(str, Enum):
    MIDTERM = "midterm"   # 中期报告
    FINAL = "final"       # 后期报告


# 默认参考文献 (编译领域经典)
DEFAULT_REFS = [
    "[1] Aho A V, Lam M S, Sethi R, Ullman J D. Systems: Principles, Techniques, and Tools (2nd Edition)[M]. Addison-Wesley, 2006.",
    "[2] Cooper K D, Torczon L. Engineering a Compiler (2nd Edition)[M]. Morgan Kaufmann, 2011.",
    "[3] Lattner C, Adve V. 中间表示: A Compilation Framework for Lifelong Program Analysis & Transformation[C]. CGO 2004: 75-86.",
    "[4] Cytron R, Ferrante J, Rosen B K, et al. Efficiently Computing Static Single Assignment Form and the Control Dependence Graph[J]. ACM TOPLAS, 1991, 13(4): 451-490.",
    "[5] Lengauer T, Tarjan R E. A Fast Algorithm for Finding Dominators in a Flowgraph[J]. ACM TOPLAS, 1979, 1(1): 121-141.",
    "[6] Cousot P, Cousot R. Abstract Interpretation: A Unified Lattice Model for Static Analysis[C]. POPL 1977: 238-252.",
    "[7] Briggs P, Cooper K D, Torczon L. Improvements to Graph Coloring Register Allocation[J]. ACM TOPLAS, 1994, 16(3): 428-455.",
    "[8] Wegman M N, Zadeck F K. Constant Propagation with Conditional Branches[J]. ACM TOPLAS, 1991, 13(2): 181-210.",
    "[9] Patterson D, Waterman A. The 目标平台 Reader: An Open Architecture Atlas[M]. Strawberry Canyon, 2017.",
    "[10] LLVM Documentation: Writing an LLVM Pass[EB/OL]. https://llvm.org/docs/WritingAnLLVMPass.html.",
]


def build_report(meta: ProjectMeta,
                 report_type: ReportType = ReportType.FINAL,
                 style_config: StyleConfig = None,
                 refs: list = None,
                 acknowledgments: str = "",
                 **section_kwargs) -> Document:
    """生成完整报告文档。

    Args:
        meta: 项目元数据 (课程/导师/学生等)
        report_type: MIDTERM (中期) 或 FINAL (后期)
        style_config: 样式配置 (None=默认学术规范)
        refs: 参考文献列表
        acknowledgments: 致谢内容 (仅后期报告)
        **section_kwargs: 传递给各章节生成函数的额外参数

    Returns:
        python-docx Document 对象 (可 .save("xxx.docx"))
    """
    if style_config is None:
        style_config = StyleConfig(
            header_text=f"{meta.college} 综合设计报告"
        )

    doc = new_document(style_config)

    # 封面 + 摘要 + 目录
    create_cover(doc, meta, style_config)
    insert_toc_field(doc)

    if report_type == ReportType.MIDTERM:
        _build_midterm(doc, meta, section_kwargs)
    else:
        _build_final(doc, meta, section_kwargs, acknowledgments)

    # 参考文献
    add_references(doc, refs or DEFAULT_REFS)

    return doc


# ═══════════════════════════════════════════════════════════════
# 中期报告: 3章
# ═══════════════════════════════════════════════════════════════

def _build_midterm(doc, meta, kwargs):
    """中期报告: 第一章 + 第二章 + 第三章。"""
    # ── 第一章 综合设计的进展情况 ──
    add_chapter_start(doc, "综合设计的进展情况", 1)

    add_heading(doc, "1.1 针对工程问题的方案设计", level=2)
    write_design_solution(doc, meta, **kwargs.get('design', {}))

    add_heading(doc, "1.2 针对工程问题的推理分析", level=2)
    write_reasoning_analysis(doc, meta, **kwargs.get('reasoning', {}))

    add_heading(doc, "1.3 针对工程问题的具体实现", level=2)
    write_implementation_progress(doc, meta, **kwargs.get('impl', {}))

    add_heading(doc, "1.4 知识技能学习情况", level=2)
    write_knowledge_skills(doc, meta, **kwargs.get('skills', {}))

    # ── 第二章 存在问题与解决方案 ──
    add_chapter_start(doc, "存在问题与解决方案", 2)
    write_problems_and_solutions(doc, meta, **kwargs.get('problems', {}))

    # ── 第三章 前期任务完成度与后续实施计划 ──
    add_chapter_start(doc, "前期任务完成度与后续实施计划", 3)
    write_completion_and_plan(doc, meta, **kwargs.get('plan', {}))


# ═══════════════════════════════════════════════════════════════
# 后期报告: 4章 + 致谢
# ═══════════════════════════════════════════════════════════════

def _build_final(doc, meta, kwargs, acknowledgments):
    """后期报告: 第一章 + 第二章 + 第三章 + 第四章 + 致谢。"""
    # ── 第一章 复杂工程问题归纳与实施方案可行性研究 ──
    add_chapter_start(doc, "复杂工程问题归纳与实施方案可行性研究", 1)

    add_heading(doc, "1.1 需求分析与建模", level=2)
    write_requirements_modeling(doc, meta, **kwargs.get('requirements', {}))

    add_heading(doc, "1.2 复杂工程问题归纳", level=2)
    write_complex_problems(doc, meta, **kwargs.get('problems_induction', {}))

    add_heading(doc, "1.3 实施方案与可行性研究", level=2)
    write_feasibility_study(doc, meta, **kwargs.get('feasibility', {}))

    # ── 第二章 存在问题与解决方案 ──
    add_chapter_start(doc, "存在问题与解决方案", 2)
    write_problems_and_solutions(doc, meta, **kwargs.get('problems', {}))

    # ── 第三章 执行情况与完成度 ──
    add_chapter_start(doc, "执行情况与完成度", 3)

    add_heading(doc, "3.1 项目执行过程概述", level=2)
    write_execution_overview(doc, meta, **kwargs.get('execution', {}))

    add_heading(doc, "3.2 核心功能完成情况", level=2)
    write_core_completion(doc, meta, **kwargs.get('core_completion', {}))

    add_heading(doc, "3.3 仿真运行效果与仿真效果", level=2)
    write_simulation_results(doc, meta, **kwargs.get('simulation', {}))

    add_heading(doc, "3.4 完成度评估", level=2)
    write_completion_evaluation(doc, meta, **kwargs.get('evaluation', {}))

    # ── 第四章 分工协作与交流情况 ──
    add_chapter_start(doc, "分工协作与交流情况", 4)
    write_team_collaboration(doc, meta, **kwargs.get('team', {}))

    # ── 致谢 ──
    if acknowledgments:
        doc.add_page_break()
        add_heading(doc, "致  谢", level=1)
        for para in acknowledgments.split('\n'):
            para = para.strip()
            if para:
                add_body_para(doc, para)
